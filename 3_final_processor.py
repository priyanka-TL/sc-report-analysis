import pandas as pd
import re
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import boto3
from dotenv import load_dotenv

load_dotenv()

# --- AWS CONFIGURATION ---
try:
    claude_client = boto3.client(
        "bedrock-runtime",
        region_name="ap-south-1",
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    )
except Exception as e:
    print(f"⚠️ AWS Client Setup Failed: {e}")
    claude_client = None

MODEL_ID = "global.anthropic.claude-sonnet-4-5-20250929-v1:0"

THEME_KNOWLEDGE_BASE = """
1. Poverty and Economic Barriers
2. Legal Document-linked Barriers
3. Child Marriage Issue
4. Distance and Accessibility Issues
5. Parental Attitudes & Socio-Cultural
6. School Infrastructure & Facility
7. Teacher Capacity & Quality
8. Safety Issues
9. Substance Abuse & Addiction
10. Other Factors
"""

def categorize_environment_aggressive(text):
    """Ultra-Aggressive Environment Classification to minimize Unmapped tags."""
    text_lower = str(text).lower()
    
    school_kw = ['school', 'teacher', 'classroom', 'class', 'student', 'education', 'study', 'teaching', 'academic', 'admission', 'enroll', 'attendance', 'grade', 'subject', 'exam', 'books', 'uniform', 'midday meal', 'mid day', 'scholarship', 'library', 'playground', 'infrastructure', 'facility', 'toilet', 'water', 'building']
    home_kw = ['parent', 'family', 'mother', 'father', 'home', 'household', 'house', 'sibling', 'brother', 'sister', 'domestic', 'child labour', 'work at home', 'income', 'alcoholic', 'migration', 'marriage', 'dowry', 'attitude', 'mindset', 'belief', 'cultural', 'discrimination']
    comm_kw = ['village', 'community', 'society', 'road', 'transport', 'bus', 'distance', 'far', 'path', 'route', 'weather', 'rain', 'heat', 'flood', 'surroundings', 'neighborhood', 'area', 'locality', 'safety', 'harassment', 'molestation', 'social pressure', 'caste', 'tribe', 'practice']
    
    s_score = sum(2 if kw in text_lower else 0 for kw in school_kw)
    h_score = sum(2 if kw in text_lower else 0 for kw in home_kw)
    c_score = sum(2 if kw in text_lower else 0 for kw in comm_kw)
    
    # Contextual boosts
    if any(kw in text_lower for kw in ['to school', 'reach school', 'go to school']): c_score += 3
    if any(kw in text_lower for kw in ['at home', 'in family', 'parent awareness']): h_score += 3
    if any(kw in text_lower for kw in ['in school', 'at school', 'lacks']): s_score += 3
    
    scores = {'School': s_score, 'Home': h_score, 'Community': c_score}
    if max(scores.values()) == 0:
        if any(w in text_lower for w in ['poor', 'poverty', 'money', 'financial']): return 'Home'
        return 'Community' # Default fallback
    return max(scores, key=scores.get)

def categorize_agency(text):
    """Classifies the driver of the solution."""
    text_lower = str(text).lower()
    comm_kw = ['community', 'together', 'collective', 'meena manch', 'chaupal', 'village', 'we will', 'committee', 'panchayat']
    ind_kw = ['parent', 'family', 'individual', 'we should', 'people should', 'personally', 'mother', 'father']
    inst_kw = ['government', 'school', 'ngo', 'administration', 'authority', 'provide', 'officer', 'department', 'teacher']
    
    scores = {
        'Community-led': sum(1 for kw in comm_kw if kw in text_lower),
        'Individual-led': sum(1 for kw in ind_kw if kw in text_lower),
        'Institutional': sum(1 for kw in inst_kw if kw in text_lower)
    }
    return max(scores, key=scores.get) if max(scores.values()) > 0 else 'Community-led'

# --- 2. FORMATTING UTILITIES ---

def set_cell_background(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def normalize_text(text):
    if pd.isna(text): return "Uncategorized"
    return re.sub(r'^\d+[\.\)\s-]*', '', str(text)).strip()

def is_valid_solution(text):
    """Filters out problem statements disguised as solutions."""
    text = str(text).lower().strip()
    problem_starters = ['lack of', 'no ', 'not enough', 'poor ', 'insufficient', 'scarcity', 'absence', 'shortage', 'due to', 'because of']
    if any(text.startswith(p) for p in problem_starters):
        return False
    return True

def clean_theme_name(text):
    """Cleans theme names, handling combined themes and empty values."""
    if pd.isna(text) or str(text).strip() == "" or str(text).lower() == "nan":
        return "Other Factors"
    
    text = str(text).strip()
    
    # Handle combined themes (e.g., "Theme A + Theme B") - Take the first one
    if '+' in text:
        text = text.split('+')[0].strip()
        
    # Remove leading numbers/bullets (e.g., "1. Poverty")
    text = re.sub(r'^\d+[\.\)\s-]*', '', text).strip()
    
    return text

THEME_KNOWLEDGE_BASE = """
1. Poverty and Economic Barriers
2. Legal Document-linked Barriers
3. Child Marriage Issue
4. Distance and Accessibility Issues
5. Parental Attitudes & Socio-Cultural
6. School Infrastructure & Facility
7. Teacher Capacity & Quality
8. Safety Issues
9. Substance Abuse & Addiction
10. Other Factors
"""

def refine_concepts_with_ai(concepts_list, type_label):
    """
    Uses AI to clean, deduplicate, and re-theme the top concepts.
    Returns a dictionary: { 'Old Concept': {'concept': 'New Concept', 'theme': 'New Theme'} }
    """
    if not claude_client: return {}
    
    print(f"   🧠 AI Refinement: Optimizing top {len(concepts_list)} {type_label}s...")
    
    all_results = {}
    batch_size = 50
    
    for i in range(0, len(concepts_list), batch_size):
        batch = concepts_list[i:i+batch_size]
        print(f"      Processing batch {i//batch_size + 1} ({len(batch)} items)...")
        
        prompt = f"""You are a Data Cleaning Expert for an Education Report.
        
        THEMES:
        {THEME_KNOWLEDGE_BASE}
        
        INPUT: A list of top recurring {type_label}s found in the data.
        
        TASKS:
        1. AGGRESSIVE DEDUPLICATION: Merge specific variants into broader core concepts.
           - "Child labor in agriculture" / "Child labor at home" / "Child labour due to poverty" / "Child labour preventing education" -> MERGE ALL INTO "Child Labour"
           - "Poverty preventing girls' education" / "Poverty preventing school attendance" / "Poverty preventing children's education" -> MERGE ALL INTO "Poverty preventing education"
           - "Lack of awareness" / "General awareness" -> MERGE INTO "Lack of awareness about education importance"
        2. RE-THEME: Correct misclassified items.
        3. FORMAT: Ensure the concept is a clear, concise {type_label} statement.
        
        INPUT LIST:
        {json.dumps(batch)}
        
        OUTPUT:
        Return a VALID JSON object where keys are the INPUT strings and values are objects with "concept" and "theme".
        IMPORTANT: 
        - Escape all double quotes within strings (e.g., \"text\").
        - Do not include any text outside the JSON block.
        - Ensure the JSON is valid.
        
        Example:
        {{
            "Child labor in agriculture": {{"concept": "Child Labour", "theme": "Poverty and Economic Barriers"}},
            "General awareness": {{"concept": "Lack of awareness about education importance", "theme": "Other Factors"}}
        }}
        RETURN ONLY JSON. NO MARKDOWN."""

        try:
            response = claude_client.invoke_model(
                modelId=MODEL_ID,
                body=json.dumps({
                    "anthropic_version": "bedrock-2023-05-31",
                    "max_tokens": 4000,
                    "messages": [{"role": "user", "content": prompt}]
                })
            )
            resp_body = json.loads(response['body'].read())
            text = resp_body['content'][0]['text'].strip()
            
            # Robust JSON extraction
            json_match = re.search(r'\{.*\}', text, re.DOTALL)
            if json_match:
                text = json_match.group(0)
            else:
                text = text.replace('```json', '').replace('```', '').strip()
                
            batch_result = json.loads(text)
            all_results.update(batch_result)
            
        except Exception as e:
            print(f"      ⚠️ Batch {i//batch_size + 1} Failed: {e}")
            
    return all_results

# --- MAIN ENGINE ---

def generate_report():
    print("🚀 Starting Final Report Generation Engine...")
    
    # 1. LOAD DATASETS
    print("   📂 Loading datasets...")
    try:
        df_raw = pd.read_csv('cleaned_data.csv') 
        chal_exploded = pd.read_csv('exploded_challenges.csv')
        sol_exploded = pd.read_csv('exploded_solutions.csv')
        chal_map = pd.read_csv('challenge_mapping.csv')
        sol_map = pd.read_csv('solution_mapping.csv')
    except Exception as e:
        print(f"❌ Error: Required CSV files missing. {e}")
        return

    # Normalize mappings
    print("   ⚙️  Processing data and applying categories...")
    chal_map['Theme'] = chal_map['Theme'].apply(clean_theme_name)
    sol_map['Theme'] = sol_map['Theme'].apply(clean_theme_name)

    # CREATE df_c (Challenges) and apply environment logic
    df_c = chal_exploded.merge(chal_map, left_on='Challenges', right_on='Original', how='left')
    df_c['Theme'] = df_c['Theme'].fillna("Other Factors").astype(str)
    df_c['Theme'] = df_c['Theme'].apply(clean_theme_name) # Double check after merge
    df_c['Environment'] = df_c['Challenges'].apply(categorize_environment_aggressive)
    
    # RETAIN original variable name for District Profile section
    df_chal_mapped = df_c 

    # CREATE df_s (Solutions) and apply agency logic
    df_s = sol_exploded.merge(sol_map, left_on='Solutions', right_on='Original', how='left')
    df_s['Theme'] = df_s['Theme'].fillna("Other Factors").astype(str)
    df_s['Theme'] = df_s['Theme'].apply(clean_theme_name) # Double check after merge
    df_s['Agency'] = df_s['Solutions'].apply(categorize_agency)

    # --- AI REFINEMENT STEP ---
    # Get top 200 concepts to refine (Increased from 100 to catch more variations)
    top_chal = df_c['Merged_Concept'].value_counts().head(200).index.tolist()
    top_sol = df_s['Merged_Concept'].value_counts().head(200).index.tolist()
    
    # Refine Challenges
    chal_updates = refine_concepts_with_ai(top_chal, "Challenge")
    if chal_updates:
        # Apply updates
        for old, new_data in chal_updates.items():
            mask = df_c['Merged_Concept'] == old
            df_c.loc[mask, 'Merged_Concept'] = new_data['concept']
            df_c.loc[mask, 'Theme'] = new_data['theme']
            
    # Refine Solutions
    sol_updates = refine_concepts_with_ai(top_sol, "Solution")
    if sol_updates:
        for old, new_data in sol_updates.items():
            mask = df_s['Merged_Concept'] == old
            df_s.loc[mask, 'Merged_Concept'] = new_data['concept']
            df_s.loc[mask, 'Theme'] = new_data['theme']

    # Re-clean themes just in case AI returned something weird
    df_c['Theme'] = df_c['Theme'].apply(clean_theme_name)
    df_s['Theme'] = df_s['Theme'].apply(clean_theme_name)
    
    # Update df_chal_mapped reference
    df_chal_mapped = df_c 

    # --- BASELINE METRIC CALCULATIONS ---
    TOTAL_CH_STATE = len(df_raw)
    for col in ['Participant Count', 'Men', 'Women', 'Children']:
        df_raw[col] = pd.to_numeric(df_raw[col], errors='coerce').fillna(0)
    
    TOTAL_PART_STATE = int(df_raw['Participant Count'].sum())
    NUM_CHAL_STATEMENTS = len(chal_exploded) # Used for % calculations
    NUM_SOL_STATEMENTS = len(sol_exploded)   # Used for % calculations
    NUM_CHAL = NUM_CHAL_STATEMENTS
    NUM_SOL = NUM_SOL_STATEMENTS
    
    SOL_RATIO = (NUM_SOL / NUM_CHAL) if NUM_CHAL > 0 else 0
    
    # Demographics
    m_total = int(df_raw['Men'].sum())
    w_total = int(df_raw['Women'].sum())
    c_total = int(df_raw['Children'].sum())
    # Calculate Others
    df_raw['Others'] = df_raw['Participant Count'] - (df_raw['Men'] + df_raw['Women'] + df_raw['Children'])
    df_raw['Others'] = df_raw['Others'].clip(lower=0) 
    o_total = int(df_raw['Others'].sum())
    
    # Percentages
    w_perc = (w_total / TOTAL_PART_STATE * 100) if TOTAL_PART_STATE > 0 else 0
    m_perc = (m_total / TOTAL_PART_STATE * 100) if TOTAL_PART_STATE > 0 else 0
    c_perc = (c_total / TOTAL_PART_STATE * 100) if TOTAL_PART_STATE > 0 else 0
    o_perc = (o_total / TOTAL_PART_STATE * 100) if TOTAL_PART_STATE > 0 else 0
    
    avg_per_chaupal = TOTAL_PART_STATE / TOTAL_CH_STATE if TOTAL_CH_STATE > 0 else 0
    num_districts = df_raw['District'].nunique()
    num_themes = df_c['Theme'].nunique()

    # Theme Analysis for Summary
    theme_counts = df_c['Theme'].value_counts()
    top_3_themes = theme_counts.head(3)
    top_3_perc = (top_3_themes.sum() / NUM_CHAL * 100) if NUM_CHAL > 0 else 0

    # Agency Analysis for Summary
    agency_counts = df_s['Agency'].value_counts()
    ind_led = agency_counts.get('Individual-led', 0)
    comm_led = agency_counts.get('Community-led', 0)
    inst_led = agency_counts.get('Institutional', 0)
    
    ind_perc = (ind_led / NUM_SOL * 100) if NUM_SOL > 0 else 0
    comm_perc = (comm_led / NUM_SOL * 100) if NUM_SOL > 0 else 0
    inst_perc = (inst_led / NUM_SOL * 100) if NUM_SOL > 0 else 0
    
    comm_driven_perc = ind_perc + comm_perc

    doc = Document()

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    print("   📝 Generating Section 1: Executive Summary...")
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
    
    # Intro Paragraph
    intro_p = doc.add_paragraph()
    intro_p.add_run(f"This comprehensive report presents an in-depth analysis of {TOTAL_CH_STATE:,} Shiksha Chaupal community dialogues conducted across Bihar, representing the collective voices of {TOTAL_PART_STATE:,} community members. These dialogues constitute one of the most extensive participatory consultations on education challenges in India, providing rich insights into grassroots barriers to education and community-driven solutions. The analysis encompasses {NUM_CHAL:,} individual challenges and {NUM_SOL:,} solutions, systematically categorized into {num_themes} primary thematic areas for comprehensive understanding.")

    # KEY INSIGHT
    doc.add_heading(f'🔑 KEY INSIGHT: Solution Coverage Ratio', level=2)
    p_ratio = doc.add_paragraph()
    run_ratio = p_ratio.add_run(f"Solution-to-Challenge Ratio: {SOL_RATIO:.2f}")
    run_ratio.bold = True
    
    if SOL_RATIO >= 1.0:
        ratio_text = f"This remarkable ratio demonstrates that communities identified {NUM_SOL:,} solutions for {NUM_CHAL:,} challenges. This transcends traditional deficit-based consultations where communities merely list problems. Instead, it reveals communities as active problem-solvers who think constructively about actionable interventions. This represents a paradigm shift in community engagement from problem identification to solution co-creation."
    elif SOL_RATIO >= 0.5:
        ratio_text = f"With {NUM_SOL:,} solutions proposed for {NUM_CHAL:,} challenges, communities are actively engaging in problem-solving. This indicates a constructive approach where participants are moving beyond just listing problems to identifying potential interventions."
    else:
        ratio_text = f"Communities identified {NUM_SOL:,} solutions alongside {NUM_CHAL:,} challenges. While the focus remains on highlighting barriers, there is an emerging capacity for solution-finding that can be further nurtured."
    
    doc.add_paragraph(ratio_text)

    # SCALE OF PARTICIPATION
    doc.add_heading('Scale of Community Participation', level=2)
    scale_p = doc.add_paragraph()
    scale_p.add_run(f"• Geographic Reach: {TOTAL_CH_STATE:,} community dialogues conducted across {num_districts} districts in Bihar\n")
    scale_p.add_run(f"• Total Participants: {TOTAL_PART_STATE:,} community members actively engaged\n")
    scale_p.add_run(f"• Average Engagement: {avg_per_chaupal:.1f} participants per Chaupal, indicating strong community mobilization\n")
    scale_p.add_run(f"• Gender Representation: Women {w_perc:.1f}%, Children {c_perc:.1f}%, Men {m_perc:.1f}%, Others {o_perc:.1f}%\n")
    
    if w_perc > 50:
        gender_text = f"• Dominant female participation ({w_perc:.1f}%) signals authentic grassroots engagement rather than tokenistic consultation, as women are primary stakeholders in children's education"
    elif w_perc > m_perc:
        gender_text = f"• Strong female participation ({w_perc:.1f}%) highlights women's active role in discussing education challenges, outnumbering male participants ({m_perc:.1f}%)."
    else:
        gender_text = f"• The dialogues included diverse participation, with women contributing {w_perc:.1f}% of the voices, ensuring maternal perspectives are included."
    
    scale_p.add_run(gender_text)

    # DOMINANT CHALLENGE THEMES
    doc.add_heading('Dominant Challenge Themes', level=2)
    doc.add_paragraph(f"The thematic analysis reveals systemic patterns in education barriers. The top 3 themes collectively account for {top_3_perc:.1f}% of all challenges, indicating concentrated problem areas requiring prioritized intervention:")
    
    for theme, count in top_3_themes.items():
        t_perc = (count / NUM_CHAL * 100)
        # Get solution count for this theme
        s_count = len(df_s[df_s['Theme'] == theme])
        doc.add_paragraph(f"{theme}: {t_perc:.1f}% ({count:,} challenges, {s_count:,} solutions)", style='List Bullet')

    # COMMUNITY-LED VS SYSTEM-DEPENDENT
    doc.add_heading('Community-Led vs System-Dependent Solutions', level=2)
    doc.add_paragraph("Solution agency analysis reveals community ownership patterns. The distribution demonstrates where communities see themselves as agents of change versus where they require external institutional support:")
    
    agency_p = doc.add_paragraph()
    agency_p.add_run(f"• Individual-led Solutions: {ind_perc:.1f}% ({ind_led:,} solutions) - Family-level actions including parental engagement, behavioral change, and household resource allocation\n")
    agency_p.add_run(f"• Community-led Solutions: {comm_perc:.1f}% ({comm_led:,} solutions) - Collective action including social mobilization, peer support networks, and community organizing\n")
    agency_p.add_run(f"• Institutional Solutions: {inst_perc:.1f}% ({inst_led:,} solutions) - Systemic interventions requiring government or CSO support including infrastructure, policy changes, and resource provision\n")
    
    if comm_driven_perc > 50:
        insight_text = f"• Critical Insight: {comm_driven_perc:.1f}% of solutions are community-driven (individual + community-led), demonstrating extraordinary grassroots capacity that partnerships must amplify rather than replace"
    else:
        insight_text = f"• Critical Insight: While {comm_driven_perc:.1f}% of solutions are community-driven, a significant portion ({inst_perc:.1f}%) requires institutional support, highlighting the need for strong government-community collaboration."

    crit_run = agency_p.add_run(insight_text)
    crit_run.bold = True

    # STRATEGIC PARTNERSHIP
    doc.add_heading('🤝 Strategic Partnership Opportunity', level=2)
    
    if comm_driven_perc > 50:
        strat_text = f"The {comm_driven_perc:.1f}% proportion of community-led and individual-led solutions reveals extraordinary community ownership and problem-solving capacity. Strategic partnerships should operate on a community-strengthening model rather than community-replacing model. This means: (1) Amplifying existing community initiatives through capacity building and resource support, (2) Providing targeted institutional interventions ({inst_perc:.1f}%) for infrastructure, teacher capacity, and documentation systems that communities cannot address independently, (3) Facilitating community-to-community learning and peer exchange, (4) Advocating for policy changes that enable community solutions to scale. The partnership must recognize communities as co-creators and primary implementers, not merely beneficiaries."
    else:
        strat_text = f"With {inst_perc:.1f}% of solutions requiring institutional intervention, a collaborative partnership model is essential. This involves: (1) Government and CSOs addressing structural barriers like infrastructure and teacher shortages, (2) Strengthening the {comm_driven_perc:.1f}% of community-led initiatives to ensure sustainability, (3) Creating feedback loops where community needs directly inform policy implementation."

    doc.add_paragraph(strat_text)

    doc.add_page_break()

    # --- SECTION 2: GENERAL PARTICIPATION OVERVIEW ---
    print("   📝 Generating Section 2: Participation Overview...")
    doc.add_heading('2. GENERAL PARTICIPATION OVERVIEW', level=1)
    
    intro_para = doc.add_paragraph()
    run = intro_para.add_run("This section provides comprehensive analysis of participation patterns across geographic and demographic dimensions.")
    run.italic = True

    # TABLE 1: STATE METRICS
    doc.add_heading('TABLE 1: Overall Participation Metrics', level=2)
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    for i, h in enumerate(['Metric', 'Count', 'Percentage (%)']):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_background(hdr_cells[i], "D9D9D9") 

    def add_row_v3(t, metric, count, p_type):
        r = t.add_row().cells
        r[0].text = metric
        r[1].text = f"{int(count):,}" if isinstance(count, (int, float)) else str(count)
        if p_type == "none": r[2].text = "-"
        elif p_type == "full": r[2].text = "100%"
        else:
            p = (count / TOTAL_PART_STATE * 100) if TOTAL_PART_STATE > 0 else 0
            r[2].text = f"{p:.1f}%"

    add_row_v3(table1, "Total Number of Chaupals", TOTAL_CH_STATE, "none")
    add_row_v3(table1, "Total Number of Participants", TOTAL_PART_STATE, "full")
    add_row_v3(table1, "Average Participants per Chaupal", f"{avg_per_chaupal:.1f}", "none")
    add_row_v3(table1, "Men Participants", m_total, "calc")
    add_row_v3(table1, "Women Participants", w_total, "calc")
    add_row_v3(table1, "Children Participants", c_total, "calc")
    add_row_v3(table1, "Others (Unspecified)", o_total, "calc")

    # SECTION 2.3: DISTRICT DISTRIBUTION
    doc.add_heading('2.3 District-wise Distribution of Reported Chaupals', level=2)
    dist_stats = df_raw.groupby('District').agg(Ch_Count=('id', 'count'), Part_Sum=('Participant Count', 'sum')).reset_index()
    dist_stats['Ch_Perc'] = (dist_stats['Ch_Count'] / TOTAL_CH_STATE) * 100
    dist_stats = dist_stats.sort_values('Ch_Count', ascending=False)

    table_dist = doc.add_table(rows=1, cols=5)
    table_dist.style = 'Table Grid'
    d_hdr = table_dist.rows[0].cells
    headers = ['District', 'Chaupals (N)', 'Chaupal %', 'Participants (N)', 'Participant %']
    for i, h in enumerate(headers):
        d_hdr[i].text = h
        d_hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_background(d_hdr[i], "F2F2F2")

    for _, row in dist_stats.iterrows():
        r = table_dist.add_row().cells
        r[0].text, r[1].text = str(row['District']), f"{int(row['Ch_Count']):,}"
        r[2].text = f"{row['Ch_Perc']:.1f}%"
        r[3].text, r[4].text = f"{int(row['Part_Sum']):,}", f"{(row['Part_Sum']/TOTAL_PART_STATE)*100:.1f}%"

    # NARRATIVE ANALYSIS
    top_3 = dist_stats.head(3)
    top_3_sum = top_3['Ch_Perc'].sum()
    dist_list_str = ", ".join([f"{row['District']} ({int(row['Ch_Count']):,}, {row['Ch_Perc']:.1f}%)" for _, row in top_3.iterrows()])
    
    doc.add_heading('Geographic Distribution Analysis', level=3)
    doc.add_paragraph(f"Geographic distribution shows concentration in {dist_list_str}. Together, these top 3 districts account for {top_3_sum:.1f}% of dialogues.")

    doc.add_heading('Demographic Composition Analysis', level=3)
    w_perc, m_perc, c_perc = (w_total/TOTAL_PART_STATE)*100, (m_total/TOTAL_PART_STATE)*100, (c_total/TOTAL_PART_STATE)*100
    ratio = w_perc / m_perc if m_perc > 0 else 0
    doc.add_paragraph(f"Women constitute {w_perc:.1f}% of participants, which is {( 'more than triple' if ratio >= 3 else 'significantly higher than' )} male participation ({m_perc:.1f}%).")

    doc.add_page_break()
    

    # --- SECTION 3: CORE CONTENT ANALYSIS ---
    print("   📝 Generating Section 3: Core Content Analysis...")
    doc.add_heading('3. CORE CONTENT ANALYSIS', level=1)
    
    # Calculations for Section 3
    unique_chal_count = chal_map['Merged_Concept'].nunique()
    unique_sol_count = sol_map['Merged_Concept'].nunique()
    chal_reduction = ((NUM_CHAL - unique_chal_count) / NUM_CHAL * 100) if NUM_CHAL > 0 else 0
    sol_reduction = ((NUM_SOL - unique_sol_count) / NUM_SOL * 100) if NUM_SOL > 0 else 0

    # Narrative
    doc.add_paragraph(f"This section analyzes the substance of community dialogues - the challenges identified and solutions proposed. Communities articulated {NUM_CHAL:,} individual challenges and {NUM_SOL:,} individual solutions.")

    # TABLE 2
    doc.add_heading('Overall Challenge & Solution Metrics', level=2)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    hdr = summary_table.rows[0].cells
    hdr[0].text, hdr[1].text = 'Metric', 'Count'
    for cell in hdr: 
        set_cell_background(cell, "D9D9D9")
        cell.paragraphs[0].runs[0].bold = True

    data_rows = [
        ("Total Challenges", f"{NUM_CHAL:,}"),
        ("Total Solutions", f"{NUM_SOL:,}"),
        ("Unique Challenges (after deduplication)", f"{unique_chal_count:,}"),
        ("Unique Solutions (after deduplication)", f"{unique_sol_count:,}"),
        ("Overall Solution-to-Challenge Ratio", f"{SOL_RATIO:.2f}"),
    ]
    for metric, val in data_rows:
        r = summary_table.add_row().cells
        r[0].text, r[1].text = metric, str(val)

    # 3.2 District Averages
    doc.add_heading('District-wise Engagement Depth', level=2)
    dist_chal_counts = chal_exploded.groupby('District').size().reset_index(name='C_Count')
    dist_sol_counts = sol_exploded.groupby('District').size().reset_index(name='S_Count')
    
    avg_df = dist_stats[['District', 'Ch_Count']].merge(dist_chal_counts, on='District').merge(dist_sol_counts, on='District')
    avg_df['Avg_C'] = avg_df['C_Count'] / avg_df['Ch_Count']
    avg_df['Avg_S'] = avg_df['S_Count'] / avg_df['Ch_Count']

    avg_table = doc.add_table(rows=1, cols=3)
    avg_table.style = 'Table Grid'
    a_hdr = avg_table.rows[0].cells
    a_hdr[0].text, a_hdr[1].text, a_hdr[2].text = 'District', 'Avg Challenges/Chaupal', 'Avg Solutions/Chaupal'
    for cell in a_hdr: set_cell_background(cell, "F2F2F2")

    for _, row in avg_df.sort_values('Avg_C', ascending=False).iterrows():
        r = avg_table.add_row().cells
        r[0].text, r[1].text, r[2].text = str(row['District']), f"{row['Avg_C']:.2f}", f"{row['Avg_S']:.2f}"

    # Paradigm Shift Section
    doc.add_heading('Solution-to-Challenge Ratio: A Paradigm Shift', level=2)
    doc.add_paragraph(f"The overall solution-to-challenge ratio of {SOL_RATIO:.2f} represents a fundamental paradigm shift in community consultation methodology. Traditional deficit-based consultations focus solely on problem identification, treating communities as problem containers. The Shiksha Chaupal model demonstrates that when communities are engaged as problem-solvers rather than merely problem-identifiers, they actively think constructively about solutions. This {NUM_SOL:,} solutions for {NUM_CHAL:,} challenges ratio indicates that every articulated challenge was matched with actionable solution thinking, demonstrating community agency and constructive engagement. This asset-based approach recognizes communities as repositories of contextual knowledge and innovative problem-solving capacity.")

    doc.add_page_break()

    # --- SECTION 4: THEMATIC ANALYSIS ---
    print("   📝 Generating Section 4: Thematic Analysis...")
    doc.add_heading('4. THEMATIC ANALYSIS', level=1)
    
    # TABLE: Challenge by Theme
    doc.add_heading('Overall Challenge Distribution by Theme', level=2)
    t_chal = doc.add_table(rows=1, cols=3); t_chal.style = 'Table Grid'
    h_chal = t_chal.rows[0].cells
    h_chal[0].text, h_chal[1].text, h_chal[2].text = 'Theme', 'Count', '%'
    for c in h_chal: set_cell_background(c, "D9D9D9")
    
    theme_counts = df_c['Theme'].value_counts()
    for theme, count in theme_counts.items():
        r = t_chal.add_row().cells
        r[0].text, r[1].text, r[2].text = theme, str(count), f"{(count/NUM_CHAL_STATEMENTS*100):.1f}%"

    # TABLE: Solution Agency
    doc.add_heading('Solution Distribution by Agency', level=2)
    t_agency = doc.add_table(rows=1, cols=3); t_agency.style = 'Table Grid'
    h_age = t_agency.rows[0].cells
    h_age[0].text, h_age[1].text, h_age[2].text = 'Agency Type', 'Count', '%'
    for c in h_age: set_cell_background(c, "F2F2F2")
    
    agency_counts = df_s['Agency'].value_counts()
    for agency, count in agency_counts.items():
        r = t_agency.add_row().cells
        r[0].text, r[1].text, r[2].text = agency, str(count), f"{(count/NUM_SOL_STATEMENTS*100):.1f}%"

    # TABLE: Challenge Environment
    doc.add_heading('Challenge Distribution by Environment', level=2)
    t_env = doc.add_table(rows=1, cols=3); t_env.style = 'Table Grid'
    h_env = t_env.rows[0].cells
    h_env[0].text, h_env[1].text, h_env[2].text = 'Environment', 'Count', '%'
    for c in h_env: set_cell_background(c, "D9D9D9")
    
    env_counts = df_c['Environment'].value_counts()
    for env, count in env_counts.items():
        r = t_env.add_row().cells
        r[0].text, r[1].text, r[2].text = env, str(count), f"{(count/NUM_CHAL_STATEMENTS*100):.1f}%"

    # --- INDIVIDUAL THEME DEEP-DIVES ---
    doc.add_page_break()

    # Prepare theme list: Top themes, but ensure "Other Factors" is last
    themes_to_process = list(theme_counts.index[:10])
    
    # If "Other Factors" is in the list, remove it temporarily
    if "Other Factors" in themes_to_process:
        themes_to_process.remove("Other Factors")
        
    # If "Other Factors" exists in data (even if not in top 10 originally, though unlikely given 20%), add it to end
    if "Other Factors" in theme_counts.index:
        themes_to_process.append("Other Factors")

    for i, theme in enumerate(themes_to_process, 1): 
        doc.add_heading(f'4.{i} {theme.upper()}', level=2)
        
        t_c = df_c[df_c['Theme'] == theme]
        t_s = df_s[df_s['Theme'] == theme]
        
        # Theme Metrics
        c_count = len(t_c)
        s_count = len(t_s)
        c_perc = (c_count / NUM_CHAL * 100) if NUM_CHAL > 0 else 0
        sol_cov = (s_count / c_count) if c_count > 0 else 0
        u_c = t_c['Merged_Concept'].nunique()
        u_s = t_s['Merged_Concept'].nunique()
        
        # Add Metrics Paragraph
        m_para = doc.add_paragraph()
        m_para.add_run(f"Scale: {c_count:,} challenges ({c_perc:.1f}% of total dataset) | {s_count:,} solutions\n").bold = True
        m_para.add_run(f"Solution Coverage: {sol_cov:.2f} solutions per challenge")

        # Challenge Landscape
        doc.add_heading('Challenge Landscape', level=4)
        if not t_c.empty:
            env_pref = t_c['Environment'].value_counts(normalize=True).idxmax()
            doc.add_paragraph(f"The landscape for '{theme}' is primarily localized within the {env_pref} environment. This suggests that interventions must be targeted at this level for maximum impact.")
        
        doc.add_heading("Top Recurring Challenges", level=5)
        
        # Logic for 50% coverage
        chal_counts = t_c['Merged_Concept'].value_counts()
        total_theme_chal = len(t_c)
        cumulative_count = 0
        
        for i, (concept, count) in enumerate(chal_counts.items(), 1):
            cumulative_count += count
            coverage_perc = (cumulative_count / total_theme_chal) * 100
            item_perc = (count / total_theme_chal) * 100
            
            # Find representative quote (longest original text for this concept)
            original_texts = t_c[t_c['Merged_Concept'] == concept]['Challenges'].tolist()
            rep_quote = max(original_texts, key=len) if original_texts else concept
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.add_run(f"{i}. {concept}").bold = True
            p.add_run(f" ({count} mentions, {item_perc:.1f}%)")
            p.add_run(f"\n   Voice from the ground: \"{rep_quote}\"").italic = True
            
            if coverage_perc >= 50 and i >= 5: # At least 5 items, or until 50%
                break
            if i >= 15: # Hard cap
                break

        # Solution Ecosystem
        doc.add_heading('Solution Ecosystem', level=4)
        if not t_s.empty:
            total_theme_sol = len(t_s)
            agency_counts = t_s['Agency'].value_counts(normalize=True)
            agency_main = agency_counts.idxmax()
            agency_perc = agency_counts.max() * 100
            
            doc.add_paragraph(f"Communities proposed {total_theme_sol:,} solutions to address this theme. The solution ecosystem demonstrates {agency_main} agency with {agency_perc:.1f}% of solutions being {agency_main}.")

            doc.add_heading("Most Frequently Proposed Solutions", level=5)
            
            # Logic for 50% coverage - Solutions
            sol_counts = t_s['Merged_Concept'].value_counts()
            cumulative_count_s = 0
            printed_count = 0
            
            for concept, count in sol_counts.items():
                if not is_valid_solution(concept):
                    continue

                printed_count += 1
                cumulative_count_s += count
                coverage_perc = (cumulative_count_s / total_theme_sol) * 100
                item_perc = (count / total_theme_sol) * 100
                
                # Find representative quote
                original_texts = t_s[t_s['Merged_Concept'] == concept]['Solutions'].tolist()
                rep_quote = max(original_texts, key=len) if original_texts else concept

                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                p.paragraph_format.first_line_indent = Pt(-18)
                p.add_run(f"{printed_count}. {concept}").bold = True
                p.add_run(f" ({count} mentions, {item_perc:.1f}%)")
                p.add_run(f"\n   Community Proposal: \"{rep_quote}\"").italic = True
                
                if coverage_perc >= 50 and printed_count >= 5:
                    break
                if printed_count >= 15:
                    break

    doc.add_page_break()
    print("   📝 Generating Section 5: District Profiles...")
    doc.add_heading('5. DISTRICT PROFILES', level=1)
    
    doc.add_paragraph("This section provides comprehensive profiles for top-performing districts by Chaupal count, including quantitative metrics, thematic breakdowns, and top challenges/solutions specific to each district. These profiles enable district-specific intervention design and comparative analysis across geographies.")

    # --- DISTRICT OVERVIEW TABLE ---
    doc.add_heading('District Performance Overview Table', level=2)
    
    # Aggregate Data
    d_stats = df_raw.groupby('District').agg(
        Chaupals=('id', 'nunique'),
        Participants=('Participant Count', 'sum')
    ).reset_index()
    
    c_counts = chal_exploded.groupby('District').size().reset_index(name='Challenges')
    s_counts = sol_exploded.groupby('District').size().reset_index(name='Solutions')
    
    dist_overview = d_stats.merge(c_counts, on='District', how='left').merge(s_counts, on='District', how='left').fillna(0)
    dist_overview['Ratio'] = dist_overview['Solutions'] / dist_overview['Challenges']
    dist_overview = dist_overview.sort_values('Chaupals', ascending=False)
    
    # Create Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['District', 'Chaupals', 'Participants', 'Challenges', 'Solutions', 'Ratio']
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_background(hdr[i], "D9D9D9")
        
    for _, row in dist_overview.iterrows():
        r = table.add_row().cells
        r[0].text = str(row['District'])
        r[1].text = f"{int(row['Chaupals']):,}"
        r[2].text = f"{int(row['Participants']):,}"
        r[3].text = f"{int(row['Challenges']):,}"
        r[4].text = f"{int(row['Solutions']):,}"
        r[5].text = f"{row['Ratio']:.2f}"
        
    # Note
    note_p = doc.add_paragraph()
    note_run = note_p.add_run("Note: Ratio = Solutions ÷ Challenges. Values >1.0 indicate more solutions than challenges identified.")
    note_run.italic = True
    note_run.font.size = Pt(9)

    # --- DETAILED PROFILES ---
    # Reorder districts to put 'Others' last
    dist_list = dist_overview['District'].tolist()
    if 'Others' in dist_list:
        dist_list.remove('Others')
        dist_list.append('Others')

    for i, dist in enumerate(dist_list, 1):
        doc.add_heading(f'5.{i} {dist.upper()}', level=2)
        d_raw = df_raw[df_raw['District'] == dist]
        
        # Filter data for this district
        d_chal = df_c[df_c['District'] == dist]
        d_sol = df_s[df_s['District'] == dist]

        # Metrics for Snapshot
        dist_chaupals = len(d_raw)
        dist_participants = int(d_raw['Participant Count'].sum())
        chaupal_perc = (dist_chaupals / TOTAL_CH_STATE * 100) if TOTAL_CH_STATE > 0 else 0
        
        m_d = d_raw['Men'].sum()
        w_d = d_raw['Women'].sum()
        c_d = d_raw['Children'].sum()
        
        m_p = (m_d / dist_participants * 100) if dist_participants > 0 else 0
        w_p = (w_d / dist_participants * 100) if dist_participants > 0 else 0
        c_p = (c_d / dist_participants * 100) if dist_participants > 0 else 0
        
        dist_chal_count = len(d_chal)
        dist_sol_count = len(d_sol)
        dist_ratio = (dist_sol_count / dist_chal_count) if dist_chal_count > 0 else 0
        
        avg_c = dist_chal_count / dist_chaupals if dist_chaupals > 0 else 0
        avg_s = dist_sol_count / dist_chaupals if dist_chaupals > 0 else 0

        doc.add_heading('A. Quantitative Snapshot', level=3)
        snap_p = doc.add_paragraph()
        snap_p.add_run(f"Chaupals: {dist_chaupals:,} ({chaupal_perc:.1f}% of total)\n")
        snap_p.add_run(f"Total Participants: {dist_participants:,}\n")
        snap_p.add_run(f"Demographics: Men {m_p:.1f}%, Women {w_p:.1f}%, Children {c_p:.1f}%\n")
        snap_p.add_run(f"Challenges: {dist_chal_count:,} | Solutions: {dist_sol_count:,}\n")
        snap_p.add_run(f"Solution Efficiency: {dist_ratio:.2f}\n")
        snap_p.add_run(f"Average per Chaupal: {avg_c:.1f} challenges, {avg_s:.1f} solutions")
        
        if d_chal.empty:
            doc.add_paragraph("No challenge data available for this district.")
            continue

        # Calculate Theme Percentages
        total_dist_chal = len(d_chal)
        theme_counts = d_chal['Theme'].value_counts()
        
        doc.add_heading('Thematic Breakdown & Examples', level=3)
        
        for theme, count in theme_counts.items():
            perc = (count / total_dist_chal) * 100
            
            # Theme Header
            p_theme = doc.add_paragraph()
            p_theme.paragraph_format.space_before = Pt(6)
            run = p_theme.add_run(f"• {theme} ({perc:.1f}%)")
            run.bold = True
            
            # Get Top 2 Challenges (by frequency in this district)
            theme_c_rows = d_chal[d_chal['Theme'] == theme]
            top_challenges = theme_c_rows['Merged_Concept'].value_counts().head(2).index.tolist()
            
            # Get Top 2 Solutions
            theme_s_rows = d_sol[d_sol['Theme'] == theme]
            
            # DEBUG
            # print(f"Theme: {theme}, theme_s_rows shape: {theme_s_rows.shape}")
            
            if theme_s_rows.empty:
                top_solutions = []
            else:
                # Ensure Merged_Concept is string
                theme_s_rows = theme_s_rows.copy()
                theme_s_rows['Merged_Concept'] = theme_s_rows['Merged_Concept'].fillna("Uncategorized").astype(str)
                
                valid_mask = theme_s_rows['Merged_Concept'].apply(is_valid_solution)
                valid_s_rows = theme_s_rows[valid_mask]
                top_solutions = valid_s_rows['Merged_Concept'].value_counts().head(2).index.tolist()
            
            # Write Challenges
            if top_challenges:
                p_c = doc.add_paragraph()
                p_c.paragraph_format.left_indent = Pt(18)
                p_c.add_run("Challenges: ").bold = True
                p_c.add_run("; ".join(top_challenges))
                
            # Write Solutions
            if top_solutions:
                p_s = doc.add_paragraph()
                p_s.paragraph_format.left_indent = Pt(18)
                p_s.add_run("Solutions: ").bold = True
                p_s.add_run("; ".join(top_solutions))

    # --- SECTION 6: UNIQUE INSIGHTS ---
    print("   📝 Generating Section 6: Unique Insights...")
    doc.add_page_break()
    doc.add_heading('6. UNIQUE INSIGHTS', level=1)
    
    def get_unique_examples(agency_type, count=5):
        subset = df_s[df_s['Agency'] == agency_type].copy()
        if subset.empty: return []
        
        # 1. Prioritize 'Other Factors'
        others = subset[subset['Theme'] == 'Other Factors']
        
        # 2. If not enough, look for low frequency items in general
        # Calculate frequency of Merged_Concept
        freq = subset['Merged_Concept'].value_counts()
        unique_concepts = freq[freq == 1].index.tolist()
        
        # Filter subset for these unique concepts
        unique_rows = subset[subset['Merged_Concept'].isin(unique_concepts)]
        
        # Combine: Others first, then unique rows
        candidates = pd.concat([others, unique_rows]).drop_duplicates(subset=['Solutions'])
        
        # Filter out short/junk text
        candidates = candidates[candidates['Solutions'].str.len() > 20]
        
        # Sort by length (longer is usually better for "insights")
        candidates['len'] = candidates['Solutions'].str.len()
        top_candidates = candidates.sort_values('len', ascending=False).head(count)
        
        results = []
        for _, row in top_candidates.iterrows():
            dist = str(row['District'])
            sol = row['Solutions'].strip()
            
            # Only show District as per request
            location_str = f"{dist}"
                
            results.append(f"\"{sol}\"\n   📍 {location_str}")
        return results

    # 6.1 Individual-led
    doc.add_heading('Individual-led Unique Solutions', level=2)
    doc.add_paragraph("Highlights of innovative or distinct solutions proposed by individuals that stand out from common themes:")
    for ex in get_unique_examples('Individual-led', 5):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(ex)

    # 6.2 Community-led
    doc.add_heading('Community-led Unique Solutions', level=2)
    doc.add_paragraph("Highlights of collective actions or community-driven innovations:")
    for ex in get_unique_examples('Community-led', 5):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(ex)

    # 6.3 Expectations
    doc.add_heading('Expectations from Government & CSOs', level=2)
    doc.add_paragraph("Key expectations and demands expressed by the community for systemic support:")
    
    inst_subset = df_s[df_s['Agency'] == 'Institutional']
    if not inst_subset.empty:
        # Get top concepts
        top_inst = inst_subset['Merged_Concept'].value_counts().head(5)
        for concept, count in top_inst.items():
            # Get a representative quote
            quotes = inst_subset[inst_subset['Merged_Concept'] == concept]['Solutions'].tolist()
            # Filter for valid quotes
            quotes = [q for q in quotes if isinstance(q, str) and len(q) > 10]
            if not quotes: continue
            
            best_quote = max(quotes, key=len)
            
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{concept}").bold = True
            p.add_run(f" ({count} mentions)")
            p.add_run(f"\n   Community Voice: \"{best_quote}\"").italic = True

    print("   💾 Saving document...")
    
    # --- SECTION 7: CONCLUSION ---
    print("   📝 Generating Section 7: Conclusion...")
    doc.add_page_break()
    doc.add_heading('7. CONCLUSION', level=1)
    
    # 7.1 Transformative Insight
    doc.add_heading('🎯 TRANSFORMATIVE INSIGHT: Community as Solution Architects', level=2)
    doc.add_paragraph(f"This analysis of {TOTAL_CH_STATE:,} Shiksha Chaupals reveals a fundamental truth: rural communities are not passive recipients of development interventions but sophisticated problem-solvers capable of designing culturally-appropriate, sustainable solutions to education barriers. The {SOL_RATIO:.2f} solution coverage ratio demonstrates that communities are generating solutions at nearly the same rate as identifying challenges.")

    # 7.2 Interconnected Challenge Reality
    doc.add_heading('The Interconnected Challenge Reality', level=2)
    top_themes_list = top_3_themes.index.tolist()
    theme_str = ", ".join(top_themes_list)
    doc.add_paragraph(f"Challenges do not exist in isolation. The dominance of themes like {theme_str} suggests a complex interplay of factors. For instance, economic barriers often amplify documentation issues, while infrastructure gaps can drive families toward alternative schooling options. Understanding these interconnections is essential for designing effective interventions.")

    # 7.3 Community Agency Excellence
    doc.add_heading('Community Agency Excellence', level=2)
    # Get top community solutions for examples
    comm_sols = df_s[df_s['Agency'] == 'Community-led']['Merged_Concept'].value_counts().head(4).index.tolist()
    if comm_sols:
        comm_examples = ", ".join([s.lower() for s in comm_sols])
        doc.add_paragraph(f"Communities have demonstrated remarkable innovation through initiatives such as {comm_examples}. This agency must be recognized, celebrated, and supported—not replaced by external solutions.")
    else:
        doc.add_paragraph("Communities have demonstrated remarkable innovation through collective action and peer support mechanisms. This agency must be recognized, celebrated, and supported—not replaced by external solutions.")

    # 7.4 Systemic Support Imperative
    doc.add_heading('Systemic Support Imperative', level=2)
    # Get top institutional solutions for examples
    inst_sols = df_s[df_s['Agency'] == 'Institutional']['Merged_Concept'].value_counts().head(4).index.tolist()
    if inst_sols:
        inst_examples = ", ".join([s.lower() for s in inst_sols])
        doc.add_paragraph(f"While community agency is exceptional, certain barriers require institutional action. Issues such as {inst_examples} cannot be resolved through community effort alone. The path forward requires strategic partnerships that amplify community strengths while providing systemic support.")
    else:
        doc.add_paragraph("While community agency is exceptional, certain barriers require institutional action. Infrastructure gaps, documentation bottlenecks, and resource shortages cannot be resolved through community effort alone. The path forward requires strategic partnerships that amplify community strengths while providing systemic support.")

    # 7.5 Strategic Partnership Opportunity
    doc.add_heading('🤝 Strategic Partnership Opportunity', level=2)
    doc.add_paragraph("The optimal collaboration model combines community-led cultural change initiatives with institutional support for infrastructure and policy barriers. Government agencies, NGOs, and CSOs should position themselves as resource partners and accountability allies—not solution designers—enabling communities to scale their own innovations while addressing systemic gaps.")

    doc.save('Final_Shiksha_Report.docx')
    print("\n🏁 SUCCESS! Complete report generated.")

if __name__ == "__main__":
    generate_report()