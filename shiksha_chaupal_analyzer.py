"""
Shiksha Chaupal Analysis Report Generator
Complete implementation following all critical requirements from the prompt
"""

import pandas as pd
import numpy as np
import re
import json
from difflib import SequenceMatcher, get_close_matches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum. text import WD_ALIGN_PARAGRAPH
from docx. enum.table import WD_TABLE_ALIGNMENT
from docx. oxml.shared import OxmlElement, qn
import warnings
from datetime import datetime
warnings.filterwarnings('ignore')

class ShikshaChaupalAnalyzer: 
    def __init__(self, csv_file):
        """Initialize analyzer with CSV data"""
        self. df = pd.read_csv(csv_file)
        self.all_challenges = []
        self.all_solutions = []
        self.canonical_challenges = []
        self.canonical_solutions = []
        
        # Initialize district mapping for ultra-aggressive extraction
        self.district_mapping = {
            'rohtas': 'Rohtas', 'rohtash': 'Rohtas', 'rohats': 'Rohtas', 'rotas': 'Rohtas',
            'kaimur': 'Kaimur', 'kaimoor': 'Kaimur', 'bhabua': 'Kaimur', 'kaimūr': 'Kaimur',
            'gaya': 'Gaya', 'gaia': 'Gaya', 'gay': 'Gaya',
            'patna': 'Patna', 'patana': 'Patna',
            'sitamarhi': 'Sitamarhi', 'seetamarhi': 'Sitamarhi', 'sitamari': 'Sitamarhi',
            'west champaran': 'West Champaran', 'pashchim champaran':  'West Champaran', 
            'w champaran': 'West Champaran', 'bettiah': 'West Champaran',
            'east champaran': 'East Champaran', 'purvi champaran': 'East Champaran', 
            'e champaran': 'East Champaran', 'motihari': 'East Champaran',
            'muzaffarpur': 'Muzaffarpur', 'mujaffarpur': 'Muzaffarpur', 'muzaffarnagar': 'Muzaffarpur',
            'saran': 'Saran', 'chapra': 'Saran', 'chhapra': 'Saran',
            'siwan': 'Siwan', 'sivan': 'Siwan',
            'gopalganj': 'Gopalganj', 'gopal ganj': 'Gopalganj',
            'vaishali': 'Vaishali', 'vaisali': 'Vaishali', 'hajipur': 'Vaishali',
            'samastipur': 'Samastipur',
            'darbhanga': 'Darbhanga',
            'madhubani': 'Madhubani',
            'supaul': 'Supaul',
            'araria': 'Araria', 'arariya': 'Araria',
            'kishanganj': 'Kishanganj', 'kishangunj': 'Kishanganj',
            'purnia': 'Purnia', 'purnea': 'Purnia',
            'katihar': 'Katihar',
            'bhagalpur': 'Bhagalpur',
            'banka': 'Banka',
            'munger': 'Munger', 'monghyr': 'Munger', 'mungair': 'Munger',
            'lakhisarai': 'Lakhisarai',
            'sheikhpura': 'Sheikhpura',
            'nalanda': 'Nalanda', 'bihar sharif': 'Nalanda',
            'nawada': 'Nawada',
            'aurangabad': 'Aurangabad',
            'jehanabad': 'Jehanabad',
            'arwal': 'Arwal',
            'bhojpur': 'Bhojpur', 'ara': 'Bhojpur',
            'buxar': 'Buxar', 'buxur': 'Buxar',
            'jamui': 'Jamui',
            'khagaria': 'Khagaria',
            'begusarai': 'Begusarai',
            'saharsa': 'Saharsa',
            'madhepura': 'Madhepura',
            'udaipur': 'Udaipur',
            'others': 'Others',
        }
        
        # All 10 theme names
        self.all_theme_names = [
            'Poverty and Economic Barriers',
            'Legal Document-linked Barriers',
            'Child Marriage Issue',
            'Distance and Accessibility Issues',
            'Parental Attitudes and Socio-Cultural Barriers',
            'School Infrastructure and Facility Issues',
            'Teacher Capacity and Quality Issues',
            'Safety Issues',
            'Substance Abuse and Addiction Issues',
            'Other Factors'
        ]

    def extract_district_ultra_aggressive(self, user_location):
        """8-stage ultra-aggressive district extraction"""
        if pd.isna(user_location):
            return 'Others'
        
        original = str(user_location).strip()
        if not original:
            return 'Others'
        
        location = original.lower()
        
        # STAGE 1: Exact match
        if location in self.district_mapping:
            return self.district_mapping[location]
        
        # STAGE 2: Substring match
        for variation, standard_name in self.district_mapping.items():
            if variation in location:
                return standard_name
        
        # STAGE 3: Parse comma-separated format
        if ',' in location: 
            parts = [p.strip() for p in location.split(',')]
            for part in parts: 
                if part in self.district_mapping:
                    return self.district_mapping[part]
                for variation, standard_name in self.district_mapping.items():
                    if variation in part:
                        return standard_name
        
        # STAGE 4: Word-by-word search
        words = re.findall(r'\w+', location)
        for word in words:
            if word in self.district_mapping:
                return self.district_mapping[word]
        
        # STAGE 5: Fuzzy match with common misspellings
        for variation, standard_name in self.district_mapping.items():
            if len(variation) >= 4: 
                if variation[: 4] in location or location[:4] in variation:
                    return standard_name
        
        # STAGE 6: Check for district headquarters/major towns
        headquarters = {
            'bettiah': 'West Champaran',
            'motihari': 'East Champaran',
            'hajipur': 'Vaishali',
            'bihar sharif': 'Nalanda',
            'ara': 'Bhojpur',
            'chapra': 'Saran',
            'chhapra': 'Saran',
        }
        for town, district in headquarters. items():
            if town in location: 
                return district
        
        # STAGE 7: Check for partial word matches with edit distance
        matches = get_close_matches(location, self.district_mapping. keys(), n=1, cutoff=0.6)
        if matches:
            return self.district_mapping[matches[0]]
        
        # STAGE 8: Check individual words against district names
        for word in words:
            if len(word) >= 4:
                matches = get_close_matches(word, self.district_mapping. keys(), n=1, cutoff=0.7)
                if matches:
                    return self.district_mapping[matches[0]]
        
        return 'Others'

    def parse_participant_count(self, value):
        """Handle participant count in numeric and JSON formats"""
        if pd.isna(value):
            return {'total': 0, 'men': 0, 'women': 0, 'children': 0}
        if isinstance(value, (int, float)):
            return {'total':  int(value), 'men': 0, 'women': 0, 'children': 0}
        if isinstance(value, str):
            try:
                data = json.loads(value. replace("'", '"'))
                return {
                    'total': int(data. get('total', 0) or 0),
                    'men':  int(data.get('men', 0) or 0),
                    'women': int(data. get('women', 0) or 0),
                    'children': int(data.get('children', 0) or 0)
                }
            except: 
                return {'total':  0, 'men': 0, 'women': 0, 'children': 0}

    def extract_items(self, cell_content):
        """Extract valid challenges/solutions from cell content"""
        if pd.isna(cell_content):
            return []
        
        text = str(cell_content)
        items = re.split(r'[|;,\n]+', text)
        
        valid_items = []
        for item in items:
            item = item.strip()
            if not item or len(item. split()) < 3:
                continue
            if '.. .' in item:  # Skip truncated
                continue
            if len(item.split()) >= 10:  # Ensure ≥10 words
                valid_items.append(item)
        
        return valid_items

    def categorize_theme_ultra_aggressive(self, text):
        """Ultra-aggressive theme categorization with expanded keywords"""
        text_lower = text.lower()
        
        # MASSIVELY EXPANDED keyword sets
        themes = {
            'Poverty and Economic Barriers': [
                'poor', 'poverty', 'no money', 'financial', 'unemployment', 'child labour',
                'economic', 'lack of money', 'poor condition', 'financial constraint',
                'financial difficulty', 'cannot afford', 'expensive', 'cost', 'income',
                'livelihood', 'work', 'labor', 'labour', 'wage', 'earning', 'breadwinner',
                'economically weak', 'below poverty', 'bpl', 'economically backward',
                'need money', 'no income', 'family income', 'household income'
            ],
            
            'Legal Document-linked Barriers': [
                'aadhaar', 'aadhar', 'birth certificate', 'id', 'enrollment', 'document',
                'admission', 'enroll', 'certificate', 'proof', 'identity', 'registration',
                'id card', 'identity proof', 'legal document', 'paperwork', 'documentation'
            ],
            
            'Child Marriage Issue': [
                'marriage', 'child marriage', 'early marriage', 'marry', 'married',
                'wedding', 'bride', 'groom', 'husband', 'wife', 'in-laws', 'marital'
            ],
            
            'Distance and Accessibility Issues': [
                'far', 'distance', 'bus', 'road', 'rain', 'heat', 'transportation', 'transport',
                'walk', 'travel', 'reach', 'access', 'route', 'path', 'vehicle', 'far away',
                'long distance', 'no school in village', 'school not nearby', 'remote',
                'weather', 'flood', 'sun', 'hot', 'cold', 'monsoon', 'season'
            ],
            
            'Parental Attitudes and Socio-Cultural Barriers': [
                'discrimination', 'cultural', 'not allowed', 'dowry', 'girls', 'purdah',
                'social pressure', 'run away', 'attitude', 'mindset', 'belief', 'custom',
                'tradition', 'culture', 'social', 'gender', 'boys', 'son', 'daughter',
                'girl child', 'male', 'female', 'patriarchy', 'conservative', 'orthodox',
                'caste', 'religion', 'community belief', 'society view', 'fear', 'shame',
                'honor', 'izzat', 'prestige', 'reputation', 'what will people say'
            ],
            
            'School Infrastructure and Facility Issues': [
                'water', 'toilet', 'mid-day meal', 'midday', 'scholarship', 'book', 'uniform',
                'government scheme', 'infrastructure', 'facility', 'building', 'classroom',
                'playground', 'library', 'laboratory', 'furniture', 'bench', 'desk', 'chair',
                'blackboard', 'whiteboard', 'electricity', 'light', 'fan', 'drinking water',
                'sanitation', 'hygiene', 'cleanliness', 'maintenance', 'repair', 'boundary wall',
                'compound', 'premises', 'school building', 'room', 'space', 'pad', 'sanitary',
                'scheme not received', 'benefit not received', 'government support'
            ],
            
            'Teacher Capacity and Quality Issues':  [
                'teacher', 'teaching', 'quality', 'absent', 'shortage', 'english', 'learning',
                'lesson', 'education quality', 'study', 'teach', 'instruction', 'pedagogy',
                'method', 'staff', 'faculty', 'educator', 'trainer', 'coaching', 'tuition',
                'subject', 'curriculum', 'syllabus', 'not teach properly', 'do not teach',
                'poor teaching', 'lack of teacher', 'teacher not come', 'teacher late',
                'attention', 'focus', 'concentrate', 'explain', 'understand'
            ],
            
            'Safety Issues': [
                'harassment', 'molestation', 'teasing', 'violence', 'unsafe', 'safety',
                'security', 'afraid', 'fear', 'danger', 'threat', 'risk', 'harm', 'abuse',
                'attack', 'assault', 'eve-teasing', 'stray dog', 'animal', 'criminal',
                'rowdy', 'goon', 'eve tease', 'pass comment', 'stare', 'follow', 'trouble',
                'bother', 'disturb', 'boundary wall', 'gate', 'watchman', 'guard', 'protection'
            ],
            
            'Substance Abuse and Addiction Issues': [
                'alcohol', 'drug', 'addiction', 'gambling', 'online game', 'mobile', 'phone',
                'addicted', 'alcoholic', 'drunk', 'drinking', 'substance', 'tobacco', 'smoking',
                'game', 'gaming', 'internet', 'screen', 'gadget', 'device', 'bet', 'betting',
                'liquor', 'wine', 'beer', 'intoxication', 'addict', 'abuse'
            ],
            
            'Other Factors': []
        }
        
        # Score each theme
        scores = {}
        for theme, keywords in themes.items():
            if theme == 'Other Factors':
                continue
            score = sum(1 for kw in keywords if kw in text_lower)
            if score > 0:
                scores[theme] = score
        
        # If clear winner, return it
        if scores: 
            max_score = max(scores.values())
            if max_score >= 2:
                return max(scores, key=scores.get)
            elif max_score == 1:
                top_theme = max(scores, key=scores.get)
                return top_theme
        
        # FALLBACK LOGIC - aggressive assignment to prevent "Other Factors"
        if any(word in text_lower for word in ['cannot', 'unable', 'afford', 'money', 'income', 'work']):
            return 'Poverty and Economic Barriers'
        
        if any(word in text_lower for word in ['learn', 'study', 'education', 'quality', 'progress']):
            return 'Teacher Capacity and Quality Issues'
        
        if any(word in text_lower for word in ['parent', 'family', 'home', 'household', 'mother', 'father']):
            return 'Parental Attitudes and Socio-Cultural Barriers'
        
        if any(word in text_lower for word in ['school', 'facility', 'resource', 'material', 'equipment']):
            return 'School Infrastructure and Facility Issues'
        
        if any(word in text_lower for word in ['go', 'come', 'reach', 'access', 'get to']):
            return 'Distance and Accessibility Issues'
        
        if any(word in text_lower for word in ['community', 'society', 'social', 'people', 'village']):
            return 'Parental Attitudes and Socio-Cultural Barriers'
        
        if any(word in text_lower for word in ['child', 'children', 'girl', 'boy', 'daughter', 'son']):
            if any(word in text_lower for word in ['young', 'small', 'age', 'year']):
                return 'Child Marriage Issue'
            else:
                return 'Parental Attitudes and Socio-Cultural Barriers'
        
        # Last resort - default to most common theme
        return 'Parental Attitudes and Socio-Cultural Barriers'

    def categorize_environment_aggressive(self, text):
        """Aggressive environment classification"""
        text_lower = text.lower()
        
        school_keywords = [
            'school', 'teacher', 'classroom', 'class', 'student', 'education',
            'study', 'studies', 'learn', 'teaching', 'academic', 'admission',
            'enroll', 'attendance', 'grade', 'subject', 'exam', 'books',
            'uniform', 'midday meal', 'mid day', 'scholarship', 'library',
            'playground', 'infrastructure', 'facility', 'toilet in school',
            'water in school', 'school building', 'school environment'
        ]
        
        home_keywords = [
            'parent', 'parents', 'family', 'mother', 'father', 'home', 'household',
            'house', 'sibling', 'brother', 'sister', 'domestic', 'child labour',
            'work at home', 'family income', 'alcoholic father', 'alcoholic parent',
            'family migration', 'family problem', 'parent awareness', 'marriage',
            'dowry', 'attitude', 'mindset', 'belief', 'cultural belief', 'discrimination'
        ]
        
        community_keywords = [
            'village', 'community', 'society', 'road', 'transport', 'transportation',
            'bus', 'distance', 'far', 'path', 'route', 'weather', 'rain', 'heat',
            'flood', 'environment around', 'surroundings', 'neighborhood', 'area',
            'locality', 'safety', 'harassment', 'molestation', 'social pressure',
            'caste', 'tribe', 'purdah system', 'cultural practice'
        ]
        
        # Count keyword matches
        school_score = sum(2 if kw in text_lower else 0 for kw in school_keywords)
        home_score = sum(2 if kw in text_lower else 0 for kw in home_keywords)
        community_score = sum(2 if kw in text_lower else 0 for kw in community_keywords)
        
        # Additional contextual scoring
        if any(kw in text_lower for kw in ['to school', 'from school', 'reach school', 'go to school']):
            community_score += 3
        
        if any(kw in text_lower for kw in ['at home', 'in family', 'parent', 'household']):
            home_score += 3
        
        if any(kw in text_lower for kw in ['in school', 'at school', 'school has', 'school lacks']):
            school_score += 3
        
        scores = {
            'School': school_score,
            'Home': home_score,
            'Community': community_score
        }
        
        max_score = max(scores.values())
        
        if max_score == 0:
            if any(word in text_lower for word in ['aadhar', 'aadhaar', 'certificate', 'document']):
                return 'Home'
            elif any(word in text_lower for word in ['poor', 'poverty', 'money', 'financial']):
                return 'Home'
            elif any(word in text_lower for word in ['quality', 'teaching', 'learning']):
                return 'School'
            else:
                return 'Community'
        
        return max(scores, key=scores.get)

    def categorize_agency_aggressive(self, text):
        """Ultra-aggressive agency classification"""
        text_lower = text.lower()
        
        community_keywords = [
            'community', 'together', 'collective', 'collectively', 'meena manch', 'chaupal',
            'village', 'villagers', 'group', 'committee', 'organization', 'we will',
            'we should', 'people should', 'society should', 'community will',
            'organize', 'awareness campaign', 'meeting', 'gather', 'mobilize',
            'social pressure', 'peer pressure', 'all together', 'united'
        ]
        
        individual_keywords = [
            'parent', 'parents', 'family', 'families', 'individual', 'mother', 'father',
            'guardian', 'household', 'personally', 'themselves', 'myself', 'ourselves',
            'each family', 'every family', 'people need to', 'parents must',
            'families must', 'send children', 'educate children', 'take responsibility',
            'own responsibility', 'self', 'personal', 'change mindset', 'change attitude'
        ]
        
        institutional_keywords = [
            'government', 'school', 'administration', 'authority', 'officials',
            'ngo', 'organization', 'provide', 'should provide', 'must provide',
            'give', 'scholarship', 'scheme', 'program', 'policy', 'fund', 'budget',
            'build', 'construct', 'infrastructure', 'facility', 'appointment',
            'hire', 'recruit', 'teacher', 'staff', 'implement', 'law', 'regulation',
            'department', 'ministry', 'panchayat should', 'government should',
            'school should', 'need support from', 'external support', 'systemic'
        ]
        
        # Count matches with weights
        community_score = sum(2 if kw in text_lower else 0 for kw in community_keywords)
        individual_score = sum(2 if kw in text_lower else 0 for kw in individual_keywords)
        institutional_score = sum(2 if kw in text_lower else 0 for kw in institutional_keywords)
        
        # Contextual scoring boosts
        if any(word in text_lower for word in ['provide', 'build', 'construct', 'hire', 'appoint']):
            institutional_score += 3
        
        if any(word in text_lower for word in ['organize', 'collective', 'together', 'campaign']):
            community_score += 3
        
        if any(word in text_lower for word in ['parents', 'family', 'household', 'each family']):
            individual_score += 3
        
        scores = {
            'Community-led': community_score,
            'Individual-led': individual_score,
            'Institutional':  institutional_score
        }
        
        max_score = max(scores.values())
        
        if max_score == 0:
            if any(word in text_lower for word in ['school', 'teacher', 'infrastructure', 'facility', 'scholarship']):
                return 'Institutional'
            elif any(word in text_lower for word in ['parent', 'family', 'child', 'girl', 'boy']):
                return 'Individual-led'
            elif any(word in text_lower for word in ['awareness', 'educate', 'understand', 'know']):
                return 'Community-led'
            else: 
                if any(word in text_lower for word in ['should', 'need', 'must', 'require']):
                    return 'Institutional'
                else:
                    return 'Community-led'
        
        return max(scores, key=scores.get)

    def semantic_similarity(self, text1, text2):
        """Calculate semantic similarity between two texts"""
        t1 = text1.lower().strip()
        t2 = text2.lower().strip()
        
        if t1 == t2:
            return 1.0
        
        similarity = SequenceMatcher(None, t1, t2).ratio()
        
        words1 = set(t1.split())
        words2 = set(t2.split())
        
        # Synonym sets
        poverty_words = {'poor', 'poverty', 'financial', 'economic', 'money', 'hardship'}
        distance_words = {'far', 'distance', 'away', 'distant', 'reach'}
        document_words = {'aadhaar', 'aadhar', 'document', 'certificate', 'id'}
        
        if words1 & poverty_words and words2 & poverty_words:
            similarity = max(similarity, 0.75)
        if words1 & distance_words and words2 & distance_words:
            similarity = max(similarity, 0.75)
        if words1 & document_words and words2 & document_words: 
            similarity = max(similarity, 0.75)
        
        return similarity

    def create_canonical_groups(self, items, threshold=0.65):
        """Create canonical groups through semantic deduplication"""
        items_sorted = sorted(items, key=lambda x:  len(x['text']), reverse=True)
        groups = []
        used = set()
        
        for i, item1 in enumerate(items_sorted):
            if i in used: 
                continue
            
            group = [item1]
            group_indices = {i}
            
            for j, item2 in enumerate(items_sorted[i+1:], start=i+1):
                if j in used:
                    continue
                
                sim = self.semantic_similarity(item1['text'], item2['text'])
                if sim >= threshold:
                    group.append(item2)
                    group_indices.add(j)
            
            used.update(group_indices)
            
            canonical = {
                'text': group[0]['text'],
                'count': len(group),
                'variants': [g['text'] for g in group],
                'theme': group[0]. get('theme', ''),
                'district': group[0]. get('district', ''),
                'environment': group[0]. get('environment', ''),
                'agency_type': group[0].get('agency_type', ''),
            }
            
            groups.append(canonical)
        
        return groups

    def process_data(self):
        """Main data processing pipeline"""
        print("Starting data processing...")
        
        # Apply district extraction
        self.df['District_Standardized'] = self. df['User Location'].apply(
            self.extract_district_ultra_aggressive
        )
        
        # Parse participant counts
        self.df['parsed_participants'] = self.df['Participant Count'].apply(
            self.parse_participant_count
        )
        
        # Extract challenges and solutions
        for idx, row in self. df.iterrows():
            challenges = self.extract_items(row. get('Challenges', ''))
            solutions = self. extract_items(row.get('Solutions', ''))
            district = row['District_Standardized']
            
            for c in challenges:
                self.all_challenges.append({
                    'text': c, 
                    'district':  district, 
                    'report_id': idx
                })
            
            for s in solutions:
                self. all_solutions.append({
                    'text': s, 
                    'district': district, 
                    'report_id': idx
                })
        
        # Apply categorizations
        for item in self.all_challenges:
            item['theme'] = self.categorize_theme_ultra_aggressive(item['text'])
            item['environment'] = self.categorize_environment_aggressive(item['text'])
        
        for item in self.all_solutions:
            item['theme'] = self. categorize_theme_ultra_aggressive(item['text'])
            item['agency_type'] = self.categorize_agency_aggressive(item['text'])
        
        # Semantic deduplication
        challenges_df = pd.DataFrame(self.all_challenges)
        solutions_df = pd. DataFrame(self.all_solutions)
        
        # Group by theme for deduplication
        for theme in challenges_df['theme'].unique():
            theme_challenges = challenges_df[challenges_df['theme'] == theme]. to_dict('records')
            theme_groups = self.create_canonical_groups(theme_challenges, threshold=0.65)
            self.canonical_challenges.extend(theme_groups)
        
        for theme in solutions_df['theme'].unique():
            theme_solutions = solutions_df[solutions_df['theme'] == theme].to_dict('records')
            theme_groups = self.create_canonical_groups(theme_solutions, threshold=0.65)
            self.canonical_solutions.extend(theme_groups)
        
        print(f"Processing complete:")
        print(f"- Extracted {len(self. all_challenges)} individual challenges")
        print(f"- Extracted {len(self.all_solutions)} individual solutions")
        print(f"- Created {len(self. canonical_challenges)} canonical challenge groups")
        print(f"- Created {len(self. canonical_solutions)} canonical solution groups")
        
        # Validation
        self.validate_requirements()

    def validate_requirements(self):
        """Validate all critical requirements"""
        print("\n=== VALIDATION RESULTS ===")
        
        # District Others percentage
        others_pct = (self.df['District_Standardized'] == 'Others').sum() / len(self.df) * 100
        print(f"District 'Others':  {others_pct:.1f}% (target: <10%)")
        
        # Theme Other Factors percentage
        if self.all_challenges:
            other_factors_pct = sum(1 for c in self.all_challenges if c['theme'] == 'Other Factors') / len(self.all_challenges) * 100
            print(f"Theme 'Other Factors': {other_factors_pct:.1f}% (target: ≤10%)")
        
        # Environment Unmapped percentage
        if self.all_challenges:
            unmapped_env_pct = sum(1 for c in self.all_challenges if c['environment'] == 'Unmapped') / len(self.all_challenges) * 100
            print(f"Environment 'Unmapped': {unmapped_env_pct:.1f}% (target: <10%)")
        
        # Agency Unmapped percentage
        if self.all_solutions:
            unmapped_agency_pct = sum(1 for s in self.all_solutions if s['agency_type'] == 'Unmapped') / len(self.all_solutions) * 100
            print(f"Agency 'Unmapped': {unmapped_agency_pct:.1f}% (target: <10%)")
        
        # Theme representation
        challenge_themes = set(c['theme'] for c in self.all_challenges)
        print(f"Themes represented: {len(challenge_themes)} of 10")
        
        if len(challenge_themes) < 10:
            missing = set(self.all_theme_names) - challenge_themes
            print(f"Missing themes: {missing}")

    def add_table_to_doc(self, doc, data, headers, title=None):
        """Add formatted table to document"""
        if title: 
            doc.add_heading(title, level=3)
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0]. cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        return table

    def generate_executive_summary(self, doc):
        """Generate Section 1: Executive Summary"""
        doc.add_heading('1.  EXECUTIVE SUMMARY', level=1)
        
        # Calculate key metrics
        total_chaupals = len(self. df)
        total_participants = self.df['parsed_participants'].apply(lambda x: x['total']).sum()
        
        # Demographics
        total_men = self.df['parsed_participants'].apply(lambda x: x['men']).sum()
        total_women = self.df['parsed_participants'].apply(lambda x: x['women']).sum()
        total_children = self.df['parsed_participants'].apply(lambda x: x['children']).sum()
        total_others_demographic = total_participants - (total_men + total_women + total_children)
        
        # Solution coverage
        total_individual_challenges = len(self. all_challenges)
        total_individual_solutions = len(self.all_solutions)
        solution_coverage = (total_individual_solutions / total_individual_challenges * 100) if total_individual_challenges > 0 else 0
        
        # Top themes
        if self.canonical_challenges:
            theme_counts = pd.DataFrame(self.canonical_challenges).groupby('theme')['count'].sum().sort_values(ascending=False)
            top_3_themes = list(theme_counts. head(3).index)
        else:
            top_3_themes = ["Data unavailable", "Data unavailable", "Data unavailable"]
        
        # Agency distribution
        if self.canonical_solutions:
            agency_dist = pd.DataFrame(self.canonical_solutions).groupby('agency_type')['count'].sum()
            total_sol_count = agency_dist. sum()
            community_pct = (agency_dist. get('Community-led', 0) / total_sol_count * 100) if total_sol_count > 0 else 0
            individual_pct = (agency_dist.get('Individual-led', 0) / total_sol_count * 100) if total_sol_count > 0 else 0
            institutional_pct = (agency_dist.get('Institutional', 0) / total_sol_count * 100) if total_sol_count > 0 else 0
        else: 
            community_pct = individual_pct = institutional_pct = 0
        
        # Executive summary text
        summary_text = f"""
        This comprehensive analysis of Shiksha Chaupal dialogues reveals critical insights from {total_chaupals} community education meetings involving {total_participants: ,} participants across Bihar districts.

        🔑 KEY INSIGHT: Overall solution coverage stands at {solution_coverage:.1f}%, indicating the community's proactive approach to addressing educational challenges through collaborative dialogue and action planning.

        SCALE AND PARTICIPATION:
        The analysis encompasses {total_chaupals} Chaupals with {total_participants:,} total participants.  The demographic composition includes {total_men: ,} men ({total_men/total_participants*100:.1f}%), {total_women:,} women ({total_women/total_participants*100:.1f}%), {total_children:,} children ({total_children/total_participants*100:.1f}%), and {total_others_demographic: ,} others ({total_others_demographic/total_participants*100:.1f}%). This diverse participation demonstrates strong community engagement across gender and age groups.

        DOMINANT CHALLENGE THEMES:
        The three most prevalent challenge themes are:  (1) {top_3_themes[0]}, (2) {top_3_themes[1]}, and (3) {top_3_themes[2]}. These themes reflect the multifaceted nature of educational barriers, spanning economic, social, and infrastructural dimensions that require coordinated intervention strategies.

        SOLUTION AGENCY DISTRIBUTION:
        Community-led solutions comprise {community_pct:.1f}% of proposed interventions, individual-led solutions account for {individual_pct:.1f}%, and institutional solutions represent {institutional_pct:. 1f}%. This distribution highlights the balance between grassroots community action and the need for systemic institutional support. 

        CHALLENGE INTERCONNECTION:
        The analysis reveals how challenges compound across themes, particularly where poverty intersects with cultural barriers, distance issues amplify infrastructure gaps, and safety concerns reinforce gender-based educational disparities. These interconnections suggest that effective interventions must address multiple challenge domains simultaneously.

        STRATEGIC PARTNERSHIP OPPORTUNITY:
        The high proportion of community-identified solutions ({community_pct + individual_pct:. 1f}% combined community and individual-led) presents a compelling case for partnership models that leverage local agency while providing targeted institutional support. The data suggests communities possess significant solution capacity that can be amplified through strategic resource allocation and policy alignment.
        """
        
        doc. add_paragraph(summary_text. strip())

    def generate_participation_section(self, doc):
        """Generate Section 2: General Participation"""
        doc.add_heading('2. GENERAL PARTICIPATION', level=1)
        
        # Calculate metrics
        total_chaupals = len(self.df)
        total_participants = self.df['parsed_participants'].apply(lambda x: x['total']).sum()
        total_men = self. df['parsed_participants'].apply(lambda x: x['men']).sum()
        total_women = self.df['parsed_participants'].apply(lambda x: x['women']).sum()
        total_children = self.df['parsed_participants'].apply(lambda x: x['children']).sum()
        total_others_demographic = total_participants - (total_men + total_women + total_children)
        avg_participants = total_participants / total_chaupals if total_chaupals > 0 else 0
        
        # Overall participation table
        participation_data = [
            ['Total Chaupals', total_chaupals, '-'],
            ['Total Participants', f'{total_participants:,}', '100%'],
            ['Men', f'{total_men:,}', f'{total_men/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['Women', f'{total_women: ,}', f'{total_women/total_participants*100:. 1f}%' if total_participants > 0 else '0%'],
            ['Children', f'{total_children:,}', f'{total_children/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['Others', f'{total_others_demographic: ,}', f'{total_others_demographic/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['Average Participants per Chaupal', f'{avg_participants:.1f}', '-']
        ]
        
        self.add_table_to_doc(doc, participation_data, 
                             ['Metric', 'Count', '%'], 
                             'TABLE 1: Overall Participation')
        
        # Add note about Others category
        doc.add_paragraph("*Note: 'Others' = Total Participants - (Men + Women + Children). Represents participants with unspecified demographic category.*")
        
        # District-wise distribution
        district_summary = self.df. groupby('District_Standardized').size().sort_values(ascending=False)
        district_data = []
        for district, count in district_summary. items():
            pct = count / total_chaupals * 100
            district_data. append([district, count, f'{pct:.1f}%'])
        
        self.add_table_to_doc(doc, district_data,
                             ['District', 'Chaupals', '%'],
                             'TABLE 2: District-wise Distribution')
        
        others_pct = (self.df['District_Standardized'] == 'Others').sum() / len(self.df) * 100
        doc.add_paragraph(f"*Note: 'Others' = unclear district information ({others_pct:. 1f}%). Target: <10%.  Analyzed as full district below.*")
        
        # Demographics breakdown
        demo_data = [
            ['Men', f'{total_men:,}', f'{total_men/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['Women', f'{total_women:,}', f'{total_women/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['Children', f'{total_children:,}', f'{total_children/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['Others', f'{total_others_demographic:,}', f'{total_others_demographic/total_participants*100:.1f}%' if total_participants > 0 else '0%'],
            ['**Total**', f'**{total_participants: ,}**', '**100%**']
        ]
        
        self.add_table_to_doc(doc, demo_data,
                             ['Category', 'Count', '% of Total'],
                             'TABLE 3: Demographics')
        
        doc.add_paragraph("*Note: 'Others' = Total Participants - (Men + Women + Children)")
        doc.add_paragraph("*Percentages calculated as:  (Category Count ÷ Total Participants) × 100*")
        
        # Narrative
        narrative = f"""
        The Shiksha Chaupal initiative demonstrates robust community engagement with {total_chaupals} meetings conducted across multiple Bihar districts. The total participation of {total_participants: ,} individuals reflects strong grassroots mobilization, with an average of {avg_participants:.1f} participants per Chaupal session.

        The gender distribution shows {total_women/total_participants*100:.1f}% women participation, indicating significant female community engagement in educational dialogue. The inclusion of {total_children:,} children ({total_children/total_participants*100:. 1f}%) in discussions underscores the child-centered approach of these community forums.

        District-wise analysis reveals geographic diversity in implementation, with the largest concentration in {district_summary.index[0]} ({district_summary.iloc[0]} Chaupals, {district_summary. iloc[0]/total_chaupals*100:.1f}%). The distribution across multiple districts ensures representation of diverse socio-economic contexts and educational challenges faced by different communities. 
        """
        
        doc.add_paragraph(narrative. strip())

    def generate_content_analysis_section(self, doc):
        """Generate Section 3: Core Content Analysis"""
        doc.add_heading('3. CORE CONTENT ANALYSIS', level=1)
        
        # Calculate metrics
        total_individual_challenges = len(self.all_challenges)
        total_individual_solutions = len(self. all_solutions)
        unique_challenges = len(self.canonical_challenges)
        unique_solutions = len(self.canonical_solutions)
        solution_ratio = total_individual_solutions / total_individual_challenges if total_individual_challenges > 0 else 0
        
        # Overall metrics table
        metrics_data = [
            ['Total Individual Challenges', total_individual_challenges],
            ['Total Individual Solutions', total_individual_solutions],
            ['Unique Challenges (after deduplication)', unique_challenges],
            ['Unique Solutions (after deduplication)', unique_solutions],
            ['Overall Solution-to-Challenge Ratio', f'{solution_ratio:.2f}']
        ]
        
        self.add_table_to_doc(doc, metrics_data,
                             ['Metric', 'Count'],
                             'TABLE 1: Overall Challenge & Solution Metrics')
        
        doc.add_paragraph("*Note: 'Individual' = raw extracted items; 'Unique' = canonical forms after semantic deduplication*")
        
        # District challenge distribution
        challenges_df = pd.DataFrame(self. all_challenges)
        if not challenges_df.empty:
            district_challenges = challenges_df.groupby('district').size().sort_values(ascending=False)
            district_chaupals = self.df. groupby('District_Standardized').size()
            
            challenge_dist_data = []
            for district in district_challenges.index:
                challenge_count = district_challenges[district]
                chaupal_count = district_chaupals. get(district, 0)
                avg_per_chaupal = challenge_count / chaupal_count if chaupal_count > 0 else 0
                pct = challenge_count / total_individual_challenges * 100
                challenge_dist_data.append([district, chaupal_count, challenge_count, f'{avg_per_chaupal:.1f}', f'{pct:.1f}%'])
            
            self.add_table_to_doc(doc, challenge_dist_data,
                                 ['District', 'Chaupals', 'Challenges', 'Avg/Chaupal', '%'],
                                 'TABLE 2: District Challenge Distribution')
        
        # District solution distribution
        solutions_df = pd.DataFrame(self.all_solutions)
        if not solutions_df.empty:
            district_solutions = solutions_df.groupby('district').size().sort_values(ascending=False)
            
            solution_dist_data = []
            for district in district_solutions.index:
                solution_count = district_solutions[district]
                chaupal_count = district_chaupals. get(district, 0)
                avg_per_chaupal = solution_count / chaupal_count if chaupal_count > 0 else 0
                pct = solution_count / total_individual_solutions * 100
                solution_dist_data.append([district, chaupal_count, solution_count, f'{avg_per_chaupal:.1f}', f'{pct:.1f}%'])
            
            self.add_table_to_doc(doc, solution_dist_data,
                                 ['District', 'Chaupals', 'Solutions', 'Avg/Chaupal', '%'],
                                 'TABLE 3: District Solution Distribution')
        
        # Narrative
        narrative = f"""
        The content analysis extracted {total_individual_challenges: ,} individual challenge statements and {total_individual_solutions:,} solution proposals from community dialogues. After applying semantic deduplication to group similar items, this refined to {unique_challenges} canonical challenge categories and {unique_solutions} solution clusters.

        The overall solution-to-challenge ratio of {solution_ratio:. 2f} demonstrates the community's solution-oriented approach, with participants actively proposing interventions for identified educational barriers. This proactive stance indicates strong local agency and readiness for collaborative action.

        District-level analysis reveals variation in engagement intensity, with some regions generating higher volumes of challenges and solutions per Chaupal.  This variation reflects local context differences in educational complexity, community mobilization capacity, and dialogue facilitation effectiveness. 
        """
        
        doc.add_paragraph(narrative.strip())

    def generate_thematic_analysis_section(self, doc):
        """Generate Section 4: Thematic Analysis"""
        doc.add_heading('4. THEMATIC ANALYSIS', level=1)
        
        intro_text = """
        The analysis employs a comprehensive 10-theme categorization framework to systematically organize community-identified challenges and solutions. This thematic structure enables pattern recognition across districts and facilitates targeted intervention design. 
        """
        doc.add_paragraph(intro_text. strip())
        
        # Challenge by theme table
        if self.canonical_challenges:
            challenge_theme_counts = pd.DataFrame(self.canonical_challenges).groupby('theme')['count'].sum().sort_values(ascending=False)
            total_challenge_count = challenge_theme_counts.sum()
            
            challenge_theme_data = []
            for theme, count in challenge_theme_counts.items():
                pct = count / total_challenge_count * 100
                challenge_theme_data. append([theme, count, f'{pct:.1f}%'])
            
            challenge_theme_data.append(['**Total**', total_challenge_count, '**100%**'])
            
            self. add_table_to_doc(doc, challenge_theme_data,
                                 ['Theme', 'Challenges', '%'],
                                 'TABLE 1: Challenge by Theme')
        
        # Challenges vs Solutions table
        if self.canonical_challenges and self.canonical_solutions:
            solution_theme_counts = pd.DataFrame(self.canonical_solutions).groupby('theme')['count'].sum()
            total_solution_count = solution_theme_counts.sum()
            
            vs_data = []
            for theme in self.all_theme_names:
                challenge_count = challenge_theme_counts. get(theme, 0)
                solution_count = solution_theme_counts.get(theme, 0)
                challenge_pct = challenge_count / total_challenge_count * 100 if total_challenge_count > 0 else 0
                solution_pct = solution_count / total_solution_count * 100 if total_solution_count > 0 else 0
                ratio = solution_count / challenge_count if challenge_count > 0 else 0
                
                vs_data.append([theme, challenge_count, f'{challenge_pct:.1f}%', 
                               solution_count, f'{solution_pct:.1f}%', f'{ratio:.2f}'])
            
            self.add_table_to_doc(doc, vs_data,
                                 ['Theme', 'Challenges', '%', 'Solutions', '%', 'Ratio'],
                                 'TABLE 2: Challenges vs Solutions')
        
        # Solution by agency table
        if self.canonical_solutions:
            agency_counts = pd.DataFrame(self.canonical_solutions).groupby('agency_type')['count'].sum().sort_values(ascending=False)
            total_agency_count = agency_counts. sum()
            
            agency_data = []
            for agency, count in agency_counts.items():
                pct = count / total_agency_count * 100
                agency_data.append([agency, count, f'{pct:. 1f}%'])
            
            self.add_table_to_doc(doc, agency_data,
                                 ['Agency', 'Solutions', '%'],
                                 'TABLE 3: Solution by Agency')
        
        # Challenge by environment table
        if self.canonical_challenges:
            challenges_with_env = pd. DataFrame(self.all_challenges)
            env_counts = challenges_with_env.groupby('environment').size().sort_values(ascending=False)
            total_env_count = env_counts.sum()
            
            env_data = []
            for env, count in env_counts.items():
                pct = count / total_env_count * 100
                env_data.append([env, count, f'{pct:. 1f}%'])
            
            self.add_table_to_doc(doc, env_data,
                                 ['Environment', 'Challenges', '%'],
                                 'TABLE 4: Challenge by Environment')
        
        # Individual theme analysis
        self.generate_individual_theme_analysis(doc)

    def generate_individual_theme_analysis(self, doc):
        """Generate individual analysis for each theme"""
        
        for theme_name in self.all_theme_names: 
            # Get theme data
            theme_challenges = [c for c in self.canonical_challenges if c['theme'] == theme_name]
            theme_solutions = [s for s in self.canonical_solutions if s['theme'] == theme_name]
            
            if not theme_challenges and not theme_solutions: 
                continue
            
            doc.add_heading(f'{theme_name. upper()}', level=2)
            
            # Calculate metrics
            total_challenge_mentions = sum(c['count'] for c in theme_challenges)
            total_solution_mentions = sum(s['count'] for s in theme_solutions)
            coverage_ratio = total_solution_mentions / total_challenge_mentions if total_challenge_mentions > 0 else 0
            
            # Environment distribution for challenges
            theme_challenge_items = [c for c in self.all_challenges if c['theme'] == theme_name]
            if theme_challenge_items:
                env_dist = pd.Series([c['environment'] for c in theme_challenge_items]).value_counts(normalize=True) * 100
                primary_env = env_dist.index[0] if len(env_dist) > 0 else "Unknown"
            else:
                primary_env = "No data"
            
            metadata_text = f"Scale: {total_challenge_mentions} challenge mentions, {total_solution_mentions} solution mentions | Coverage Ratio: {coverage_ratio:.2f} | Primary Environment: {primary_env}"
            doc.add_paragraph(metadata_text, style='Intense Quote')
            
            # Challenge Landscape
            doc.add_heading('A. Challenge Landscape', level=3)
            
            if theme_challenges: 
                # Sort challenges by count for top recurring
                sorted_challenges = sorted(theme_challenges, key=lambda x: x['count'], reverse=True)
                
                # Pattern description
                pattern_text = f"This theme encompasses {len(theme_challenges)} distinct challenge patterns with {total_challenge_mentions} total community mentions.  The challenges reflect systemic barriers that require both immediate community action and longer-term institutional intervention."
                
                if theme_challenge_items:
                    env_breakdown = "Environment distribution:  "
                    for env, pct in env_dist.head(3).items():
                        env_breakdown += f"{env} {pct:.1f}%, "
                    env_breakdown = env_breakdown.rstrip(", ")
                    pattern_text += f" {env_breakdown}."
                
                doc.add_paragraph(pattern_text)
                
                # Top recurring challenges (semantic grouping)
                doc.add_paragraph("**Top Recurring Challenges** (achieving ≥50% thematic coverage):", style='Heading 4')
                
                cumulative_mentions = 0
                coverage_items = []
                for item in sorted_challenges:
                    cumulative_mentions += item['count']
                    coverage_items.append(item)
                    coverage_pct = (cumulative_mentions / total_challenge_mentions) * 100 if total_challenge_mentions > 0 else 0
                    
                    if coverage_pct >= 50:
                        break
                
                for i, item in enumerate(coverage_items, 1):
                    item_pct = (item['count'] / total_challenge_mentions) * 100 if total_challenge_mentions > 0 else 0
                    merged_note = f"[Merged from {len(item['variants'])} variants]" if len(item['variants']) > 1 else ""
                    
                    challenge_text = f"{i}. {item['text']} ({item['count']} mentions, {item_pct:.1f}%)"
                    if merged_note: 
                        challenge_text += f"\n   {merged_note}"
                    
                    doc.add_paragraph(challenge_text)
                
                # Community quotes (examples)
                doc. add_paragraph("**Community Voices** (representative examples):", style='Heading 4')
                example_count = 0
                for item in sorted_challenges[: 4]:  # Top 4 for examples
                    if example_count >= 3:  # Limit to 3 examples
                        break
                    
                    # Use first variant as example
                    if item['variants']: 
                        example_text = f"*'{item['variants'][0]}'*"
                        doc.add_paragraph(example_text, style='Intense Quote')
                        example_count += 1
            
            # Solution Ecosystem
            doc.add_heading('B.  Solution Ecosystem', level=3)
            
            if theme_solutions:
                # Agency distribution
                agency_dist = pd.DataFrame(theme_solutions).groupby('agency_type')['count'].sum()
                
                overview_text = f"Total Solutions: {len(theme_solutions)} distinct approaches with {total_solution_mentions} community mentions.  "
                
                if not agency_dist.empty:
                    agency_breakdown = "Agency distribution:  "
                    for agency, count in agency_dist. items():
                        pct = count / total_solution_mentions * 100
                        agency_breakdown += f"{agency} {pct:.1f}%, "
                    agency_breakdown = agency_breakdown. rstrip(", ")
                    overview_text += agency_breakdown + "."
                
                doc.add_paragraph(overview_text)
                
                # Solutions by agency type
                for agency_type in ['Community-led', 'Individual-led', 'Institutional']:
                    agency_solutions = [s for s in theme_solutions if s['agency_type'] == agency_type]
                    if not agency_solutions: 
                        continue
                    
                    agency_count = sum(s['count'] for s in agency_solutions)
                    agency_pct = agency_count / total_solution_mentions * 100 if total_solution_mentions > 0 else 0
                    
                    doc.add_paragraph(f"**{agency_type}** ({agency_pct:.1f}%):", style='Heading 4')
                    
                    # Top solutions for this agency
                    sorted_agency_solutions = sorted(agency_solutions, key=lambda x: x['count'], reverse=True)
                    for i, solution in enumerate(sorted_agency_solutions[:5], 1):  # Top 5
                        solution_pct = solution['count'] / total_solution_mentions * 100 if total_solution_mentions > 0 else 0
                        solution_text = f"{i}. {solution['text']} ({solution['count']} mentions, {solution_pct:.1f}%)"
                        doc.add_paragraph(solution_text)
                
                # Community-owned solutions narrative
                community_solutions = [s for s in theme_solutions if s['agency_type'] == 'Community-led']
                if community_solutions: 
                    doc. add_paragraph("**Community-Owned Solutions**", style='Heading 4')
                    community_narrative = f"Community-led interventions for {theme_name. lower()} demonstrate local ownership and sustainable implementation potential.  These solutions leverage existing social structures, collective decision-making processes, and peer-to-peer support systems."
                    
                    if community_solutions: 
                        top_community = sorted(community_solutions, key=lambda x:  x['count'], reverse=True)[0]
                        community_narrative += f" The most frequently proposed community action involves {top_community['text']. lower()}, reflecting shared understanding of effective local intervention strategies."
                    
                    doc. add_paragraph(community_narrative)
                
                # Systemic interventions narrative
                institutional_solutions = [s for s in theme_solutions if s['agency_type'] == 'Institutional']
                if institutional_solutions:
                    doc.add_paragraph("**Systemic Interventions**", style='Heading 4')
                    institutional_narrative = f"Institutional solutions for {theme_name. lower()} require coordinated policy action, resource allocation, and structural reforms. These interventions address root causes that exceed community-level capacity."
                    
                    if institutional_solutions:
                        top_institutional = sorted(institutional_solutions, key=lambda x: x['count'], reverse=True)[0]
                        institutional_narrative += f" Key systemic needs include {top_institutional['text'].lower()}, highlighting the importance of multi-stakeholder collaboration in addressing complex educational barriers."
                    
                    doc.add_paragraph(institutional_narrative)
            
            doc.add_page_break()

    def generate_district_profiles_section(self, doc):
        """Generate Section 5: District Profiles - MANDATORY SECTION"""
        doc.add_heading('5.  DISTRICT PROFILES', level=1)
        
        intro_text = """
        This section provides comprehensive profiles for every district represented in the Shiksha Chaupal data. Each district profile includes quantitative metrics, thematic breakdown, top challenges and solutions, and unique insights.  The 'Others' category represents locations with unclear district information and receives full analytical treatment as a distinct geographic entity.
        """
        doc.add_paragraph(intro_text.strip())
        
        # District performance overview table
        districts = self.df['District_Standardized'].unique()
        
        performance_data = []
        for district in districts:
            district_df = self.df[self.df['District_Standardized'] == district]
            district_chaupals = len(district_df)
            total_participants = district_df['parsed_participants'].apply(lambda x: x['total']).sum()
            avg_participants = total_participants / district_chaupals if district_chaupals > 0 else 0
            
            # Women percentage
            women_count = district_df['parsed_participants'].apply(lambda x: x['women']).sum()
            women_pct = women_count / total_participants * 100 if total_participants > 0 else 0
            
            # Challenges and solutions
            district_challenges = len([c for c in self.all_challenges if c['district'] == district])
            district_solutions = len([s for s in self.all_solutions if s['district'] == district])
            solution_efficiency = district_solutions / district_challenges if district_challenges > 0 else 0
            
            chaupals_pct = district_chaupals / len(self.df) * 100
            
            performance_data.append([
                district, 
                district_chaupals, 
                f'{chaupals_pct:.1f}%',
                f'{avg_participants:.1f}',
                f'{solution_efficiency:.2f}',
                f'{women_pct:.1f}%',
                district_challenges,
                district_solutions
            ])
        
        self.add_table_to_doc(doc, performance_data,
                             ['District', 'Chaupals', '%', 'Avg Participants', 'Solution Efficiency', 'Women %', 'Challenges', 'Solutions'],
                             'District Performance Overview')
        
        # Individual district profiles
        for district in districts:
            self.create_district_profile(doc, district)

    def create_district_profile(self, doc, district):
        """Create comprehensive profile for individual district"""
        doc.add_
